#!/usr/bin/env python3
"""
Test Script 4: File Processor
Biblioteki: openpyxl, pillow, python-docx
"""
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont
from docx import Document

print("=" * 60)
print("FILE PROCESSOR - Test openpyxl, pillow, python-docx")
print("=" * 60)

# Excel file
wb = Workbook()
ws = wb.active
ws['A1'] = "Test"
ws['B1'] = "Data"
ws.append([1, "Row 1"])
ws.append([2, "Row 2"])
wb.save('test.xlsx')
print("✓ OpenPyXL - Utworzono plik test.xlsx")

# Image processing
img = Image.new('RGB', (200, 100), color=(73, 109, 137))
d = ImageDraw.Draw(img)
d.text((10, 40), "Test Image", fill=(255, 255, 0))
img.save('test.png')
print("✓ Pillow - Utworzono obraz test.png")

# Word document
doc = Document()
doc.add_heading('Test Document', 0)
doc.add_paragraph('This is a test paragraph.')
doc.save('test.docx')
print("✓ Python-docx - Utworzono dokument test.docx")

print("\n✓ Wszystkie pliki utworzone!")
